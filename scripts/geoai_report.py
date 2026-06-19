"""geoai_report.py - Research Report Generator"""
import os
from datetime import datetime

class ResearchReport:
    def __init__(self, output_dir="report"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def generate_title(self, title, authors, affiliation, journal=None):
        s = f"# {title}\n\n"
        s += f"**Authors:** {authors}\n"
        s += f"**Affiliation:** {affiliation}\n"
        if journal: s += f"**Target Journal:** {journal}\n"
        s += f"**Date:** {datetime.now().strftime('%Y-%m-%d')}\n"
        s += "---\n\n"
        return s

    def section_introduction(self, background, research_gap, objectives):
        return f"""## 1. Introduction\n\n{background}\n\n### Research Gap\n{research_gap}\n\n### Objectives\n{objectives}\n\n"""

    def section_study_area(self, location, extent, climate, data_sources):
        return f"""## 2. Study Area\n\n**Location:** {location}\n\n**Extent:** {extent}\n\n**Climate:** {climate}\n\n**Data Sources:** {data_sources}\n\n"""

    def section_data(self, data_table, preprocessing):
        return f"""## 3. Data\n\n{data_table}\n\n### Preprocessing\n{preprocessing}\n\n"""

    def section_methods(self, methods_text, workflow_desc):
        return f"""## 4. Methods\n\n{methods_text}\n\n### Workflow\n{workflow_desc}\n\n"""

    def section_results(self, results_text, key_findings):
        return f"""## 5. Results\n\n{results_text}\n\n### Key Findings\n{key_findings}\n\n"""

    def section_discussion(self, discussion_text, limitations):
        return f"""## 6. Discussion\n\n{discussion_text}\n\n### Limitations\n{limitations}\n\n"""

    def section_conclusion(self, conclusion_text):
        return f"""## 7. Conclusion\n\n{conclusion_text}\n\n"""

    def add_figures(self, figure_paths, captions):
        s = ""
        for path, cap in zip(figure_paths, captions):
            if os.path.exists(path):
                s += f"![{cap}]({path})\n\n*{cap}*\n\n"
            else:
                s += f"*Figure: {cap}*\n\n"
        return s

    def add_references(self, refs):
        s = "## References\n\n"
        for i, ref in enumerate(refs, 1):
            s += f"[{i}] {ref}\n\n"
        return s

    def build_report(self, sections):
        report = ""
        for section in sections:
            report += section + "\n"
        return report

    def save_markdown(self, content, filename="report.md"):
        path = os.path.join(self.output_dir, filename)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return path

    def save_word(self, content, filename="report.docx"):
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            import re
            doc = Document()
            for line in content.split("\n"):
                line = line.strip()
                if line.startswith("# "):
                    doc.add_heading(line[2:], 0)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], 1)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], 2)
                elif line.startswith("**") and line.endswith("**"):
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip("*"))
                    run.bold = True
                elif line.startswith("!"):
                    continue
                elif line.strip():
                    doc.add_paragraph(line)
            path = os.path.join(self.output_dir, filename)
            doc.save(path)
            return path
        except ImportError:
            return "python-docx not installed"

    def sci_framework(self, title=""):
        return [
            "# Title\n\nAbstract here...\n\nKeywords: ...\n\n",
            "## 1. Introduction\n\n[Background, research gap, objectives]\n\n",
            "## 2. Study Area and Data\n\n[Study area description, data sources]\n\n",
            "## 3. Methods\n\n[Methodology description]\n\n",
            "## 4. Results\n\n[Results and analysis]\n\n",
            "## 5. Discussion\n\n[Interpretation and comparison]\n\n",
            "## 6. Conclusion\n\n[Summary and future work]\n\n",
            "## References\n\n[References]\n\n",
        ]

    def thesis_framework(self, level="master"):
        sections = ["# Thesis Title\n\nAbstract\n\n"]
        sections.append("## Chapter 1: Introduction\n\n")
        sections.append("## Chapter 2: Literature Review\n\n")
        sections.append("## Chapter 3: Study Area and Data\n\n")
        sections.append("## Chapter 4: Methodology\n\n")
        sections.append("## Chapter 5: Results and Analysis\n\n")
        sections.append("## Chapter 6: Discussion\n\n")
        sections.append("## Chapter 7: Conclusion\n\n")
        sections.append("## References\n\n")
        if level == "phd":
            sections.insert(4, "## Chapter 3.5: Additional Method\n\n")
        return sections


if __name__ == "__main__":
    r = ResearchReport()
    s = r.generate_title("Sample", "Author", "University")
    s += r.section_introduction("Background text", "Gap text", "Objectives")
    s += r.section_conclusion("Conclusion text")
    path = r.save_markdown(s)
    print(f"Report saved: {path}")
